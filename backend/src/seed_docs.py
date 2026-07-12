"""Synthetic textual document generator for clientpulse (File Search / RAG input).

Usage:
    uv run python -m src.seed_docs [--accounts 6] [--out data/synthetic]

Reads seeded accounts from the database and writes a per-account document set so
the text matches the structured data (ARR, term, renewal date, SLA %, sentiment,
signals). PDFs: Contract/MSA, SOW, SLA schedule (PyMuPDF target). DOCX: weekly
status update, QBR notes (python-docx target). A couple of accounts get a
deliberately missing document to exercise the "missing data lowers confidence"
path. Run `src.seed` first.
"""

from __future__ import annotations

import argparse
import datetime
import decimal
from pathlib import Path

from docx import Document
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors
from sqlalchemy import select

from .db.session import get_session_factory
from .models import Account, SlaMetric


def money(n: decimal.Decimal | float | None) -> str:
    if n is None:
        return "n/a"
    v = float(n)
    if v >= 1e6:
        m = v / 1e6
        return f"${m:.0f}M" if m % 1 == 0 else f"${m:.1f}M"
    if v >= 1e3:
        return f"${round(v / 1e3)}K"
    return f"${round(v)}"


def fdate(d: datetime.date | None) -> str:
    return d.strftime("%d %b %Y") if d else "n/a"


# --- PDF builders -----------------------------------------------------------
def _pdf(path: Path, title: str, blocks: list) -> None:
    styles = getSampleStyleSheet()
    h = ParagraphStyle("H", parent=styles["Title"], alignment=TA_CENTER, fontSize=16)
    doc = SimpleDocTemplate(str(path), pagesize=LETTER, title=title,
                            topMargin=0.8 * inch, bottomMargin=0.8 * inch)
    flow = [Paragraph(title, h), Spacer(1, 0.3 * inch)]
    for b in blocks:
        flow.append(b)
        flow.append(Spacer(1, 0.12 * inch))
    doc.build(flow)


def _para(text: str):
    return Paragraph(text, getSampleStyleSheet()["BodyText"])


def _heading(text: str):
    return Paragraph(f"<b>{text}</b>", getSampleStyleSheet()["Heading2"])


def _kv_table(rows: list[tuple[str, str]]) -> Table:
    t = Table([[k, v] for k, v in rows], colWidths=[2.4 * inch, 3.6 * inch])
    t.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#5a6577")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LINEBELOW", (0, 0), (-1, -1), 0.4, colors.HexColor("#e5e7ee")),
    ]))
    return t


def contract_pdf(path: Path, a: Account) -> None:
    auto = "auto-renews unless terminated 60 days prior" if a.contract_term_months and a.contract_term_months >= 24 else "does not auto-renew"
    _pdf(path, f"Master Services Agreement — {a.name}", [
        _para(f"This Master Services Agreement (\"MSA\") is entered into between HCLTech "
              f"(\"Provider\") and {a.name} (\"Client\"), a {a.segment} sector organization."),
        _heading("§1 Term & Renewal"),
        _kv_table([
            ("Contract term", f"{a.contract_term_months} months"),
            ("Start date", fdate(a.contract_start_dt)),
            ("Renewal date", fdate(a.renewal_dt)),
            ("Renewal", auto),
        ]),
        _heading("§4 Commercials"),
        _kv_table([
            ("Annual recurring revenue", money(a.arr)),
            ("Currency", "USD"),
            ("Payment terms", "Net-60 from invoice date"),
        ]),
        _heading("§7 Service Levels"),
        _para("Provider shall meet the service levels defined in the SLA Schedule "
              "(Exhibit A). Sustained breach of the 95% attainment target across three "
              "consecutive months entitles Client to service credits per §7.2."),
        _heading("§12 Liability"),
        _para("Provider's aggregate liability is capped at the total fees paid in the "
              "twelve (12) months preceding the claim."),
        _heading("§18 Termination & Renewal Option"),
        _para(f"Either party may terminate for material breach uncured after 30 days. "
              f"Renewal negotiations open 90 days before the renewal date ({fdate(a.renewal_dt)})."),
    ])


def sow_pdf(path: Path, a: Account) -> None:
    _pdf(path, f"Statement of Work — {a.name}", [
        _para(f"This Statement of Work (\"SOW\") is governed by the MSA between HCLTech "
              f"and {a.name}."),
        _heading("Scope"),
        _para(f"Managed delivery and support services for {a.name}'s core {a.segment.lower()} "
              f"platform, including run operations, enhancement work, and quarterly business reviews."),
        _heading("Deliverables"),
        _para("• Monthly service report &nbsp; • Quarterly business review (QBR) &nbsp; "
              "• Incident management to agreed SLAs &nbsp; • Change request handling"),
        _heading("Acceptance"),
        _para("Deliverables are accepted on written sign-off by the Client account owner "
              f"({a.owner}) within 5 business days."),
        _heading("Commercials"),
        _kv_table([
            ("Engagement value (ARR)", money(a.arr)),
            ("Term", f"{a.contract_term_months} months"),
            ("Account owner", a.owner or "n/a"),
        ]),
    ])


def sla_pdf(path: Path, a: Account, sla_rows: list[SlaMetric]) -> None:
    header = [["Period", "Target", "Actual", "Status"]]
    body = []
    for s in sla_rows:
        body.append([
            s.period_month.strftime("%b %Y"),
            f"{float(s.target_pct):.0f}%",
            f"{float(s.actual_pct):.0f}%",
            "BREACH" if s.breached else "Met",
        ])
    tbl = Table(header + body, colWidths=[1.6 * inch, 1.2 * inch, 1.2 * inch, 1.4 * inch])
    tbl.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef0f4")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#e5e7ee")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
    ]))
    breaches = sum(1 for s in sla_rows if s.breached)
    _pdf(path, f"SLA Schedule (Exhibit A) — {a.name}", [
        _para("Service level targets and recent attainment. Target attainment is 95% "
              "across all priority incidents."),
        _heading("Attainment — last 6 months"),
        tbl,
        _heading("Summary"),
        _para(f"{breaches} of {len(sla_rows)} months breached the 95% target. "
              + ("Service credits may apply per MSA §7.2." if breaches >= 3
                 else "Within contractual tolerance.")),
    ])


# --- DOCX builders ----------------------------------------------------------
def tier_of(sla_rows: list[SlaMetric], renew_days: int | None) -> str:
    """Derive a health tier from ground-truth SLA attainment + renewal proximity."""
    avg = sum(float(s.actual_pct) for s in sla_rows) / len(sla_rows) if sla_rows else 90.0
    if avg >= 95 and (renew_days is None or renew_days > 110):
        return "healthy"
    if avg >= 90:
        return "watch"
    if avg >= 84:
        return "risk"
    return "critical"


_RAG = {"healthy": "GREEN", "watch": "AMBER", "risk": "AMBER", "critical": "RED"}
_SENTIMENT = {"healthy": "Positive", "watch": "Neutral", "risk": "Mixed", "critical": "Negative"}
_PROGRESS = {
    "healthy": "All workstreams on track; client actively engaged and expanding scope.",
    "watch": "Delivery broadly on plan; a couple of minor items being worked.",
    "risk": "Delivery slipping in places; backlog growing and one escalation open.",
    "critical": "Delivery materially behind; repeated SLA misses and active client escalation.",
}
_SENT_NOTE = {
    "Positive": "Client expressed strong satisfaction with delivery and is open to expansion.",
    "Neutral": "Client is satisfied overall; no major concerns raised.",
    "Mixed": "Client raised some concerns alongside positive feedback; escalations noted.",
    "Negative": "Client expressed dissatisfaction; escalation and churn risk are elevated.",
}


def status_docx(path: Path, a: Account, tier: str, sla_rows: list[SlaMetric]) -> None:
    doc = Document()
    doc.add_heading(f"Weekly Status Update — {a.name}", level=0)
    doc.add_paragraph(f"Week ending {fdate(datetime.date.today())} · Account owner: {a.owner}")
    doc.add_paragraph(f"Overall RAG status: {_RAG[tier]}")

    doc.add_heading("Progress", level=1)
    doc.add_paragraph(_PROGRESS[tier])

    doc.add_heading("Risks & blockers", level=1)
    breaches = [s for s in sla_rows if s.breached]
    if breaches:
        months = ", ".join(s.period_month.strftime("%b") for s in breaches)
        doc.add_paragraph(f"SLA below the 95% target in: {months}.", style="List Bullet")
    if tier in ("risk", "critical"):
        doc.add_paragraph("Client has raised concern over delivery performance.", style="List Bullet")
    if not breaches and tier in ("healthy", "watch"):
        doc.add_paragraph("No material risks this period; SLAs met.", style="List Bullet")

    doc.add_heading("Next steps", level=1)
    doc.add_paragraph("Maintain delivery cadence; prep upcoming QBR.", style="List Bullet")
    doc.save(str(path))


def qbr_docx(path: Path, a: Account, tier: str) -> None:
    sentiment = _SENTIMENT[tier]
    doc = Document()
    doc.add_heading(f"QBR Notes — {a.name}", level=0)
    doc.add_paragraph(f"Quarterly Business Review · {a.segment} · Account owner: {a.owner}")

    doc.add_heading("Client sentiment", level=1)
    doc.add_paragraph(f"{sentiment}. {_SENT_NOTE[sentiment]}")

    doc.add_heading("Commercial context", level=1)
    doc.add_paragraph(f"ARR: {money(a.arr)} · Term: {a.contract_term_months} months · "
                      f"Renewal: {fdate(a.renewal_dt)}")

    doc.add_heading("Discussion", level=1)
    if tier == "healthy":
        doc.add_paragraph("Reviewed strong performance; discussed expansion opportunities.", style="List Bullet")
    elif tier == "critical":
        doc.add_paragraph("Reviewed repeated SLA failures and renewal risk; remediation plan requested.", style="List Bullet")
    else:
        doc.add_paragraph("Reviewed service performance, open risks, and renewal posture.", style="List Bullet")
    doc.save(str(path))


# --- orchestration ----------------------------------------------------------
def pick_subset(session, n: int) -> list[Account]:
    """First n accounts by uid (tier is baked into each account's SLA ground truth)."""
    accts = list(session.scalars(select(Account).order_by(Account.account_uid)))
    return accts[:n]


def main() -> None:
    ap = argparse.ArgumentParser(description="Generate synthetic textual documents.")
    ap.add_argument("--accounts", type=int, default=6, help="number of accounts to document")
    ap.add_argument("--out", default="data/synthetic", help="output directory")
    args = ap.parse_args()

    out_root = Path(args.out)
    out_root.mkdir(parents=True, exist_ok=True)

    factory = get_session_factory()
    written = 0
    with factory() as session:
        subset = pick_subset(session, args.accounts)
        today = datetime.date.today()
        for idx, a in enumerate(subset):
            sla_rows = list(session.scalars(
                select(SlaMetric).where(SlaMetric.account_uid == a.account_uid)
                .order_by(SlaMetric.period_month)
            ))
            renew_days = (a.renewal_dt - today).days if a.renewal_dt else None
            tier = tier_of(sla_rows, renew_days)

            d = out_root / a.account_uid
            d.mkdir(parents=True, exist_ok=True)

            # deliberate gaps: 1st account missing SLA schedule, 2nd missing QBR notes
            gap_sla = idx == 0
            gap_qbr = idx == 1

            contract_pdf(d / "contract_msa.pdf", a); written += 1
            sow_pdf(d / "statement_of_work.pdf", a); written += 1
            if not gap_sla:
                sla_pdf(d / "sla_schedule.pdf", a, sla_rows); written += 1
            status_docx(d / "weekly_status.docx", a, tier, sla_rows); written += 1
            if not gap_qbr:
                qbr_docx(d / "qbr_notes.docx", a, tier); written += 1

            gaps = [g for g, on in [("SLA schedule", gap_sla), ("QBR notes", gap_qbr)] if on]
            gap_note = f"  (missing: {', '.join(gaps)})" if gaps else ""
            print(f"  {a.account_uid}  {a.name[:26]:26}  {tier:8}{gap_note}")

    print(f"Wrote {written} documents for {len(subset)} accounts → {out_root}/")


if __name__ == "__main__":
    main()
